"""File reader utilities for extracting text from CV/resume documents."""

import os
import subprocess

import docx
import pdfplumber

from langchain_core.tools import tool

@tool
def extract_cv_text(file_path: str) -> str:
    """Extract text content from a CV/resume file.

    Supports PDF, DOCX, and DOC formats.
    Args:
        file_path: Path to the CV file.

    Returns:
        Extracted text content from the file, or an error message
        if extraction fails or format is unsupported.
    """
    ext = os.path.splitext(file_path)[-1].lower()
    
    if ".docx" in ext:
        try:
            doc = docx.Document(file_path)
            text = [para.text for para in doc.paragraphs]
            return '\n'.join(text)
        except Exception as e:
            return f"Error reading the .doc file: {e}"
        
    elif ".pdf" in ext:
        try:
            text = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            return '\n'.join(text)
        except Exception as e:
            return f"Error reading the .pdf file: {e}"
    
    elif ".doc" in ext:
        try:
            
            temp_docx = file_path + ".temp.docx"
            subprocess.run(['soffice', '--headless', '--convert-to', 'docx', '--outdir', os.path.dirname(file_path), file_path], check=True)
            
            doc = docx.Document(os.path.splitext(file_path)[0] + ".docx")
            text = [para.text for para in doc.paragraphs]
            
            os.remove(os.path.splitext(file_path)[0] + ".docx")
            return '\n'.join(text)
        except Exception as e:
            return f"Error reading .doc file: {e}"
    
    else:
        return "Unsupported file format! Please use .pdf, .docx, or .doc."